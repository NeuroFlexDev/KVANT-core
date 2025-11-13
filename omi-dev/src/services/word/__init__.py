from docx.oxml.ns import qn
from docx.oxml import OxmlElement, ns
from docx import Document
from docx.shared import Mm, Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from datetime import datetime

from docx.enum.style import WD_STYLE_TYPE
from io import BytesIO

class ExportWord:
    def print_header(self, cell, text):
        cell.text = text
        paragraph = cell.paragraphs[0]
        paragraph.style = "UserTableHeaderStyle"

    def print_data(self, cell, text):
        cell.text = text
        paragraph = cell.paragraphs[0]
        paragraph.style = "UserTableDataStyle"

    def create_element(self, name):
        return OxmlElement(name)

    def create_attribute(self, element, name, value):
        element.set(ns.qn(name), value)

    def add_top_header(self, document, caption, dt):
        # удаляем автоматически создающийся пустой параграф в колонтитуле
        paragraph = document.sections[0].header.paragraphs[0]
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None

        header = document.sections[0].header
        table = header.add_table(rows=1, cols=2, width=Cm(19.5))
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = caption
        hdr_cells[1].text = dt
        table.columns[0].width = Cm(15)
        table.rows[0].cells[0].width = Cm(15)
        table.columns[1].width = Cm(4.5)
        table.rows[0].cells[1].width = Cm(4.5)
        table.rows[0].cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        table.rows[0].cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        table.rows[0].cells[0].paragraphs[0].paragraph_format.space_after = Mm(0)
        table.rows[0].cells[1].paragraphs[0].paragraph_format.space_after = Mm(0)


        for row in table.rows:
            for cell in row.cells:
                cp = cell.paragraphs[0].runs
                cp[0].font.size = Pt(8)

    def add_page_number(self, paragraph):
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        page_run = paragraph.add_run()
        t1 = self.create_element('w:t')
        self.create_attribute(t1, 'xml:space', 'preserve')
        t1.text = 'Page '
        page_run._r.append(t1)

        page_num_run = paragraph.add_run()

        fldChar1 = self.create_element('w:fldChar')
        self.create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = self.create_element('w:instrText')
        self.create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar2 = self.create_element('w:fldChar')
        self.create_attribute(fldChar2, 'w:fldCharType', 'end')

        page_num_run._r.append(fldChar1)
        page_num_run._r.append(instrText)
        page_num_run._r.append(fldChar2)

        of_run = paragraph.add_run()
        t2 = self.create_element('w:t')
        self.create_attribute(t2, 'xml:space', 'preserve')
        t2.text = ' of '
        of_run._r.append(t2)

        fldChar3 = self.create_element('w:fldChar')
        self.create_attribute(fldChar3, 'w:fldCharType', 'begin')

        instrText2 = self.create_element('w:instrText')
        self.create_attribute(instrText2, 'xml:space', 'preserve')
        instrText2.text = "NUMPAGES"

        fldChar4 = self.create_element('w:fldChar')
        self.create_attribute(fldChar4, 'w:fldCharType', 'end')

        num_pages_run = paragraph.add_run()
        num_pages_run._r.append(fldChar3)
        num_pages_run._r.append(instrText2)
        num_pages_run._r.append(fldChar4)

    def omi(self, id, data, params, calc_type):
        output = BytesIO()
        document = Document()

        # Параметры страницы
        section = document.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Mm(10)
        section.right_margin = Mm(10)
        section.top_margin = Mm(10)
        section.bottom_margin = Mm(10)
        section.header_distance = Mm(5)
        section.footer_distance = Mm(5)

        # Колонтитулы
        caption = "Объект: " + params.get("a1") + " по адресу: " + params.get("a2")
        dt = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
        self.add_top_header(document, caption, dt)

        self.add_page_number(document.sections[0].footer.paragraphs[0])
        document.sections[0].different_first_page_header_footer = True

        # Титульная страница
        style = document.styles.add_style('UserStyle1', WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(30)
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        style.paragraph_format.space_before = Cm(8.5)

        style = document.styles.add_style('UserStyle2', WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(8)
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        style = document.styles.add_style('UserStyle3', WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(18)
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        style = document.styles.add_style('UserStyle4', WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(14)
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        style.paragraph_format.space_before = Cm(8)

        style = document.styles.add_style('UserTableHeaderStyle', WD_STYLE_TYPE.PARAGRAPH)
        # style.font.name = "Arial Narrow Bold"
        style.font.size = Pt(9)
        style.font.bold = True
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        style.paragraph_format.space_before = Cm(0)
        style.paragraph_format.space_after = Cm(0)

        style = document.styles.add_style('UserTableDataStyle', WD_STYLE_TYPE.PARAGRAPH)
        # style.font.name = "Arial Narrow"
        style.font.size = Pt(9)
        # style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        style.paragraph_format.space_before = Cm(0)
        style.paragraph_format.space_after = Cm(0)

        document.add_paragraph(params.get("a1"), style="UserStyle1")
        document.add_paragraph("Наименование объекта", style="UserStyle2")
        document.add_paragraph(params.get("a2"), style="UserStyle3")
        document.add_paragraph("Адрес объекта", style="UserStyle2")
        document.add_paragraph(params.get("a3"), style="UserStyle3")
        document.add_paragraph("Функциональное назначение объекта", style="UserStyle2")
        document.add_paragraph("", style="UserStyle3")
        document.add_paragraph("Расчет стоимости реализации инвестиционного объекта с помощью программного комплекса «Квант»", style="UserStyle3").italic = True

        document.add_paragraph(f"{datetime.now().year} г.", style="UserStyle4")

        # style = document.styles['UserTableStyle']
        # style.font.size = Pt(11)

        document.add_page_break()

        # Автоматическое оглавление
        document.add_paragraph("Содержание", style="UserStyle4")

        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')  # creates a new element
        fldChar.set(qn('w:dirty'), 'true')
        fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
        # change 1-3 depending on heading levels you need
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:t')
        fldChar3.text = "Right-click to update field."
        fldChar2.append(fldChar3)

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        r_element = run._r
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar4)
        p_element = paragraph._p

        document.add_page_break()

        # style = document.styles['UserTableStyle']
        # style.font.size = Pt(11)

        # Краткие характеристики и показатели по ОЛ
        summaries = data.get("summary")
        row = -1
        for summary in summaries:
            row += 1
            document.add_heading(summary.get("name"), level=3)

            if "Блок 3" not in summary.get("name"):
                table = document.add_table(rows=1, cols=4)
                # table.style = 'Medium Grid 3 Accent 5'
                hdr_cells = table.rows[0].cells
                self.print_header(hdr_cells[0], "№")
                self.print_header(hdr_cells[1], "Наименование и техническая характеристика")
                self.print_header(hdr_cells[2], "Значение параметра")
                self.print_header(hdr_cells[3], "Единица измерения")
                table.columns[0].width = Cm(1.5)
                table.rows[0].cells[0].width = Cm(1.5)
                table.columns[1].width = Cm(11.5)
                table.rows[0].cells[1].width = Cm(11.5)
                table.columns[2].width = Cm(4)
                table.rows[0].cells[2].width = Cm(4)
                table.columns[3].width = Cm(2.5)
                table.rows[0].cells[3].width = Cm(2.5)
                for element in summary.get("elements"):
                    row += 1
                    datatype = element.get("datatype")
                    if datatype == "int":
                        value = str(element.get("value")) or ""
                    elif datatype == "float":
                        value = f'{element.get("value") or 0:,.2f}'.replace(',', ' ')
                    else:
                        value = element.get("value") or ""
                    # print(value)
                    row_cells = table.add_row().cells
                    self.print_data(row_cells[0], element.get("pos"))
                    self.print_data(row_cells[1], element.get("name"))
                    self.print_data(row_cells[2], value)
                    self.print_data(row_cells[3], element.get("unit") or "")
                    row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    if element.get("params") is not None:
                        for param in element.get("params"):
                            row += 1
                            row_cells = table.add_row().cells
                            pos = param.get("pos")
                            self.print_data(row_cells[0], str(pos) if pos else "")
                            self.print_data(row_cells[1], param.get("name"))
                            self.print_data(row_cells[2], str(param.get("value")))
                            self.print_data(row_cells[3], param.get("unit"))
                            row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                            row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                table = document.add_table(rows=1, cols=4)
                hdr_cells = table.rows[0].cells
                self.print_header(hdr_cells[0], "№.")
                self.print_header(hdr_cells[1], "Наименование и техническая характеристика")
                self.print_header(hdr_cells[2], "руб")
                self.print_header(hdr_cells[3], "руб/м2")
                table.columns[0].width = Cm(1.5)
                table.rows[0].cells[0].width = Cm(1.5)
                table.columns[1].width = Cm(11.5)
                table.rows[0].cells[1].width = Cm(11.5)
                table.columns[2].width = Cm(4)
                table.rows[0].cells[2].width = Cm(4)
                table.columns[3].width = Cm(2.5)
                table.rows[0].cells[3].width = Cm(2.5)
                for element in summary.get("elements"):
                    row += 1
                    value = element.get("value")
                    if value is None:
                        if calc_type in element:
                            value = element.get(calc_type).get("value")
                    if value is not None:
                        value = f'{value:,.2f}'.replace(',', ' ')
                    else:
                        value = ""
                    value_per_unit = element.get("value_per_unit")
                    if value_per_unit is None:
                        if calc_type in element:
                            value_per_unit = element.get(calc_type).get("value_per_unit")
                    if value_per_unit is not None:
                        value_per_unit = f'{value_per_unit:,.2f}'.replace(',', ' ')
                    else:
                        value_per_unit = ""
                    # print(value_per_unit, type(value_per_unit))
                    row_cells = table.add_row().cells
                    self.print_data(row_cells[0], element.get("pos"))
                    self.print_data(row_cells[1], element.get("name"))
                    self.print_data(row_cells[2], value)
                    self.print_data(row_cells[3], value_per_unit)
                    row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Разделы ОЛ
        sections = data.get("sections")
        for section in sections:
            document.add_heading(section.get("name"), level=3)
            table = document.add_table(rows=1, cols=5)
            # table.style = 'Medium Grid 3 Accent 5'
            hdr_cells = table.rows[0].cells

            self.print_header(hdr_cells[0], "№")
            self.print_header(hdr_cells[1], "Наименование оборудования и материалов")
            self.print_header(hdr_cells[2], "Кол-во")
            self.print_header(hdr_cells[3], "Ед. изм.")
            self.print_header(hdr_cells[4], "Стоимость")

            table.columns[0].width = Cm(1.5)
            table.rows[0].cells[0].width = Cm(1.5)
            table.columns[1].width = Cm(11)
            table.rows[0].cells[1].width = Cm(11)
            table.columns[2].width = Cm(2)
            table.rows[0].cells[2].width = Cm(2)
            table.columns[3].width = Cm(2)
            table.rows[0].cells[3].width = Cm(2)
            table.columns[4].width = Cm(3)
            table.rows[0].cells[4].width = Cm(3)

            # table.rows[0].cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # table.rows[0].cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            # table.rows[0].cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # table.rows[0].cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            # table.rows[0].cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # table.rows[0].cells[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            # table.rows[0].cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # table.rows[0].cells[3].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            # table.rows[0].cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # table.rows[0].cells[4].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            for element_type in section.get("element_types"):
                hdr_cells = table.add_row().cells
                self.print_header(hdr_cells[0], element_type.get("pos"))
                self.print_header(hdr_cells[1], element_type.get("name"))
                self.print_header(hdr_cells[2], "")
                self.print_header(hdr_cells[3], "")
                self.print_header(hdr_cells[4], "")
                hdr_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                for predict in element_type.get("predicts"):
                    row_cells = table.add_row().cells
                    self.print_data(row_cells[0], predict.get("pos"))
                    self.print_data(row_cells[1], predict.get("name"))
                    self.print_data(row_cells[2], f'{predict.get("quantity"):,.0f}'.replace(',', ' '))
                    self.print_data(row_cells[3], predict.get("unit"))
                    self.print_header(row_cells[4], "")

                    row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    row_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                row_cells = table.add_row().cells
                self.print_header(row_cells[0], "")
                self.print_header(row_cells[1], "")
                self.print_header(row_cells[2], "")
                self.print_header(row_cells[3], "")
                self.print_header(row_cells[4], "")

            for key in section.get(calc_type).keys():
                hdr_cells = table.add_row().cells
                cost = section.get(calc_type).get(key)
                self.print_header(hdr_cells[0], "")
                self.print_header(hdr_cells[1], cost.get("name"))
                self.print_header(hdr_cells[2], "")
                self.print_header(hdr_cells[3], "")
                self.print_header(hdr_cells[4], f'{cost.get("value"):,.2f}'.replace(',', ' '))
                hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                hdr_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # row_cells = table.add_row().cells
        # self.print_header(row_cells[0], "")
        # self.print_header(row_cells[1], "")
        # self.print_header(row_cells[2], "")
        # self.print_header(row_cells[3], "")
        # self.print_header(row_cells[4], "")

        # total_cost = data.get("result").get("total_cost_est")
        # hdr_cells = table.add_row().cells
        # self.print_header(hdr_cells[0], "")
        # self.print_header(hdr_cells[1], "ОБЩИЙ ИТОГ")
        # self.print_header(hdr_cells[2], "")
        # self.print_header(hdr_cells[3], "")
        # self.print_header(hdr_cells[4], f'{total_cost:,.2f}'.replace(',', ' '))
        # hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # hdr_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        document.save(output)

        return output.getvalue()
